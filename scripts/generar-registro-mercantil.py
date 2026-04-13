"""
Script generador del Registro Mercantil / Acta Constitutiva
para Boogie Rent C.A. (Sociedad Anónima)
Formato legal venezolano
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Configuracion
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs", "legal")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "00-Registro-Mercantil-Boogie-Rent-CA.docx")
FONT_NAME = "Arial"
FONT_SIZE = 12
FONT_SIZE_TITLES = 14
FONT_SIZE_SUBTITLES = 12
LINE_SPACING = 1.5
MARGIN_LEFT = Cm(3)
MARGIN_RIGHT = Cm(2.5)
MARGIN_TOP = Cm(3)
MARGIN_BOTTOM = Cm(2.5)


def create_doc():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    
    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    
    return doc


def add_title(doc, text, level=0):
    """Add a title/heading to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    if level == 0:
        run.font.size = Pt(16)
        run.underline = True
    elif level == 1:
        run.font.size = Pt(FONT_SIZE_TITLES)
    else:
        run.font.size = Pt(FONT_SIZE_SUBTITLES)
    return p


def add_centered(doc, text, bold=False, size=None):
    """Add centered text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return p


def add_justified(doc, text, bold=False):
    """Add justified paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.bold = bold
    return p


def add_article(doc, title, content):
    """Add an article with title and body text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(title)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run2 = p2.add_run(content)
    run2.font.name = FONT_NAME
    run2.font.size = Pt(FONT_SIZE)
    p2.paragraph_format.first_line_indent = Cm(1.25)
    p2.paragraph_format.line_spacing = LINE_SPACING
    
    return p2


def add_article_multi(doc, title, paragraphs_content):
    """Add an article with multiple paragraphs."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(title)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    
    for content in paragraphs_content:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run2 = p2.add_run(content)
        run2.font.name = FONT_NAME
        run2.font.size = Pt(FONT_SIZE)
        p2.paragraph_format.first_line_indent = Cm(1.25)
        p2.paragraph_format.line_spacing = LINE_SPACING


def add_signature_line(doc, name, cedula, position):
    """Add a signature block."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("")
    run.font.name = FONT_NAME
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("_________________________________")
    run2.font.name = FONT_NAME
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(name)
    run3.bold = True
    run3.font.name = FONT_NAME
    
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f"V-{cedula}")
    run4.font.name = FONT_NAME
    
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run(position)
    run5.font.name = FONT_NAME
    run5.italic = True


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tcBorders>')
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>')
    tcPr.append(tcMar)


def add_inventory_table(doc, items, col_widths=None):
    if col_widths is None:
        col_widths = [Cm(1.2), Cm(4.5), Cm(1.5), Cm(8.0)]

    headers = ["N°", "DESCRIPCIÓN DEL BIEN", "CANT.", "OBSERVACIONES"]
    table = doc.add_table(rows=1 + len(items), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.width = col_widths[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        set_cell_border(cell)
        set_cell_margins(cell)

    col_alignments = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
    ]

    for i, (num, desc, cant, uni, dim, obs) in enumerate(items):
        row = table.rows[i + 1]
        desc_text = f"{desc}. {dim}" if dim != "—" else desc
        values = [num, desc_text, cant, obs]

        for j, val in enumerate(values):
            cell = row.cells[j]
            cell.width = col_widths[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = col_alignments[j]
            run = p.add_run(val)
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            set_cell_border(cell)
            set_cell_margins(cell)

    return table


def generate():
    doc = create_doc()
    
    # ==================== ENCABEZADO ====================
    add_centered(doc, "ACTA CONSTITUTIVA", bold=True, size=18)
    add_centered(doc, "Y", bold=False, size=14)
    add_centered(doc, "ESTATUTOS SOCIALES", bold=True, size=18)
    add_centered(doc, "BOOGIE RENT C.A.", bold=True, size=16)
    add_centered(doc, "SOCIEDAD ANONIMA", bold=True, size=14)
    
    doc.add_paragraph()  # Espacio
    
    add_centered(doc, "_______________", size=12)
    doc.add_paragraph()
    
    # ==================== CLAUSULA INICIAL ====================
    add_justified(doc, 
        "En la ciudad de Caracas, Distrito Capital, de la Republica Bolivariana de Venezuela, "
        "a los trece (13) dias del mes de abril del ano dos mil veintiséis (2026), se reune en "
        "esta ciudad de Caracas, Distrito Capital, ante el ciudadano Registrador Público que suscribe "
        "y en su presencia, los ciudadanos que a continuación se identifican:"
    )
    
    doc.add_paragraph()
    
    # ==================== COMPARECENCIA SOCIO 1 ====================
    add_article(doc, "PRIMERO:",
        "Sebastián Chacón Peña, venezolano, mayor de edad, de este domicilio, "
        "programador, titular de la cédula de identidad Nro. V-27.333.855, "
        "soltero, domiciliado en Caracas, Distrito Capital, Municipio Libertador, "
        "Avenida Panteón, Edf. Roda, Apto. 12-A."
    )
    
    # ==================== COMPARECENCIA SOCIO 2 ====================
    add_article(doc, "SEGUNDO:",
        "Nuryvel Antonieta Peña Gonzalez, venezolana, mayor de edad, de este domicilio, "
        "abogada, titular de la cédula de identidad Nro. V-10.505.412, casada bajo el "
        "régimen de la sociedad conyugal vigente, domiciliada en Caracas, Distrito Capital, "
        "Municipio Libertador, Avenida Panteón, Edf. Roda, Apto. 12-A."
    )
    
    doc.add_paragraph()
    
    # ==================== DECLARACION DE VOLUNTAD ====================
    add_article(doc, "TERCERO:",
        "Quienes previa las formalidades legales exigidas por el Código de Comercio de "
        "Venezuela y demas leyes aplicables, declaran su voluntad de constituir entre si "
        "una Sociedad Anónima de comercio, la cual se regirá por las disposiciones contenidas "
        "en el presente documento y se denominará con el nombre y bajo las condiciones que "
        "seguidamente se establecen:"
    )
    
    doc.add_paragraph()
    
    # ==================== NOMBRE DE LA SOCIEDAD ====================
    add_centered(doc, "DENOMINACIÓN SOCIAL", bold=True, size=14)
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA PRIMERA:",
        "La Sociedad se denomina BOOGIE RENT C.A., sociedad mercantil constituida "
        "bajo la forma de Sociedad Anónima, la cual podrá ser designada en todos sus "
        "actos y documentos legales."
    )
    
    doc.add_paragraph()
    
    # ==================== OBJETO SOCIAL ====================
    add_centered(doc, "OBJETO SOCIAL", bold=True, size=14)
    doc.add_paragraph()
    
    add_article_multi(doc, "CLÁUSULA SEGUNDA:", [
        "El objeto principal de la Sociedad será el desarrollo y explotación de actividades "
        "relacionadas con la intermediación y prestación de servicios tecnológicos, comerciales "
        "y financieros, en especial:",
        
        "1. El desarrollo, operación y explotación de plataformas tecnológicas y aplicaciones "
        "informáticas para la intermediación de servicios de alojamiento vacacional y turístico, "
        "incluyendo la gestión de reservas, procesamiento de pagos, verificación de identidad, "
        "sistema de reseñas, promociones y servicios complementarios.",
        
        "2. La prestación de servicios de tecnología de la información y comunicaciones, "
        "incluyendo desarrollo de software, aplicaciones móviles, páginas web, sistemas de "
        "información, consultoría tecnológica y servicios de alojamiento de datos.",
        
        "3. La intermediación financiera de carácter no bancario, incluyendo la gestión y "
        "procesamiento de pagos electrónicos, sistemas de monedero digital (Wallet), "
        "conversiones de divisas y servicios de pago en moneda nacional y extranjera, "
        "conforme a la legislación venezolana aplicable.",
        
        "4. La comercialización, distribución, venta al por mayor y menor, e instalación de "
        "máquinas expendedoras automáticas de productos y servicios, incluyendo su "
        "mantenimiento, operación y explotación.",
        
        "5. La comercialización al por mayor y menor de productos de consumo, higiene personal, "
        "alimentos, bebidas y bienes de consumo general a través de plataformas electrónicas "
        "y establecimientos físicos.",
        
        "6. La prestación de servicios gastronómicos, de transporte, limpieza, turismo y "
        "servicios complementarios vinculados a la industria del alojamiento vacacional.",
        
        "7. La creación, gestión y explotación de marcas, patentes, licencias, modelos "
        "de utilidad y todo tipo de propiedad intelectual, industrial y comercial.",
        
        "8. La realización de actividades publicitarias, de promoción, marketing digital "
        "y servicios de comunicación social.",
        
        "9. La adquisición, administración, enajenación y disposición de bienes muebles "
        "e inmuebles, acciones y valores mobiliarios en general.",
        
        "10. La compra, venta, permuta, arrendamiento, administración y gestión de bienes "
        "inmuebles, tanto a nombre propio de la Sociedad como en calidad de intermediaria, "
        "corredora o representante ante terceros, incluyendo la prestación de servicios de "
        "mediación inmobiliaria.",
        
        "11. La intermediación y prestación de servicios de seguros médicos temporales, "
        "de viaje y complementarios de salud, actuando como corredor o agente de seguros "
        "ante compañías aseguradoras autorizadas, conforme a la legislación venezolana "
        "aplicable.",
        
        "12. En general, la realización de cualquier actividad lícita de carácter comercial, "
        "industrial, financiero o de servicios que la Junta Directiva considere conveniente "
        "al desarrollo de la Sociedad.",
        
        "En ningún caso la Sociedad ejercerá las actividades de banco, institución financiera "
        "registrada, casa de bolsa o empresa aseguradora que actuare directamente como "
        "aseguradora, sin la autorización previa y expresa "
        "de los organismos competentes."
    ])
    
    doc.add_paragraph()
    
    # ==================== DOMICILIO ====================
    add_centered(doc, "DOMICILIO Y DURACIÓN", bold=True, size=14)
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA TERCERA:",
        "La Sociedad tendrá su domicilio principal en la ciudad de Caracas, Distrito Capital, "
        "República Bolivariana de Venezuela, específicamente en la Avenida Anauco, Edf. Golden "
        "Village, Apto. 6D, y podrá establecer sucursales, agencias, representaciones o cualquier "
        "otro tipo de establecimiento comercial en cualquier lugar del territorio nacional o del "
        "exterior que la Junta Directiva considere conveniente."
    )
    
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA CUARTA:",
        "La Sociedad tendrá una duración de CINCUENTA (50) ANOS contados a partir de su "
        "inscripción en el Registro Mercantil del lugar de su domicilio, salvo que sea prorrogada "
        "o disuelta antes por las causales previstas en la Ley."
    )
    
    doc.add_paragraph()
    
    # ==================== CAPITAL SOCIAL ====================
    add_centered(doc, "CAPITAL SOCIAL", bold=True, size=14)
    doc.add_paragraph()
    
    add_article_multi(doc, "CLÁUSULA QUINTA:", [
        "El capital social de la Sociedad se fija en la cantidad de CUARENTA MIL DOLARES "
        "DE LOS ESTADOS UNIDOS DE AMÉRICA (USD $40.000,00), o su equivalente en bolívares "
        "según la tasa oficial del Banco Central de Venezuela vigente al momento de la transacción, "
        "representado por CUARENTA MIL "
        "(40.000) acciones comunes, ordinarias y sin valor nominal, las cuales quedan "
        "distribuidas y suscritas de la siguiente manera:",
        
        "a) El ciudadano Sebastián Chacón Peña, venezolano, mayor de edad, de este domicilio, "
        "titular de la cédula de identidad Nro. V-27.333.855, suscribe TREINTA Y NUEVE MIL "
        "DOSCIENTAS (39.200) acciones comunes, lo que representa el NOVENTA Y OCHO POR CIENTO "
        "(98%) del capital social.",
        
        "b) La ciudadana Nuryvel Antonieta Peña Gonzalez, venezolana, mayor de edad, de este "
        "domicilio, titular de la cédula de identidad Nro. V-10.505.412, suscribe OCHOCIENTAS "
        "(800) acciones comunes, lo que representa el DOS POR CIENTO (2%) del capital social.",
        
        "El total de acciones suscritas asciende a CUARENTA MIL (40.000) acciones comunes, "
        "representando el CIEN POR CIENTO (100%) del capital social."
    ])
    
    doc.add_paragraph()
    
    add_article_multi(doc, "CLÁUSULA SEXTA:", [
        "Las acciones suscritas por los fundadores están pagadas en su TOTALIDAD al momento "
        "de la constitución de la Sociedad, en dólares de los Estados Unidos de América (USD), "
        "o su equivalente en bolívares según la tasa oficial del Banco Central de Venezuela vigente "
        "al momento del pago, conforme a las constancias de pago emitidas a cada suscriptor.",
        
        "Las acciones serán representadas por títulos nominativos expedidos a nombre del "
        "respectivo accionista y contendrán las menciones exigidas por el Código de Comercio. "
        "Los títulos se entregarán una vez la Sociedad quede legalmente constituida."
    ])
    
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA SÉPTIMA:",
        "La Sociedad podrá aumentar el capital social por decisión de la Asamblea de "
        "Accionistas, conforme a las formalidades establecidas en el Código de Comercio. "
        "En tal caso, los accionistas tendrán derecho preferente a suscribir nuevas acciones "
        "en proporción a las que ya posean, salvo que la Asamblea disponga otra cosa por "
        "mayoría calificada."
    )
    
    doc.add_paragraph()
    
    # ==================== INVENTARIO DE ACTIVOS ====================
    add_centered(doc, "INVENTARIO DE ACTIVOS", bold=True, size=14)
    add_centered(doc, "Constituyentes del Aporte de Capital", bold=False, size=12)
    doc.add_paragraph()
    
    add_article_multi(doc, "CLÁUSULA SÉPTIMA BIS:", [
        "Los comparecientes, en su carácter de fundadores, declaran que los activos "
        "que se detallan a continuación constituyen el aporte de capital de la Sociedad "
        "al momento de su constitución. El presente inventario tiene carácter vinculante "
        "y forma parte integral de estos Estatutos Sociales.",
        
        "Los activos quedan organizados en las siguientes categorías:"
    ])
    
    doc.add_paragraph()
    
    # ---- TABLA 1: OFICINA Y MOBILIARIO ----
    add_centered(doc, "CATEGORÍA 1: OFICINA Y MOBILIARIO", bold=True, size=12)
    doc.add_paragraph()
    
    add_justified(doc,
        "Bienes inmuebles por destino y mobiliario destinado al funcionamiento "
        "operativo y administrativo de la sede principal de la Sociedad."
    )
    doc.add_paragraph()
    
    oficina_items = [
        ("1", "Oficina principal", "1", "Unidad", "70 m²", "Oficina principal de operaciones de la empresa, ubicada en la Avenida Anauco, Edf. Golden Village, Apto. 6D, Caracas, Distrito Capital, amoblada y equipada."),
        ("2", "Escritorios ejecutivos de trabajo", "3", "Unidades", "120 x 60 cm c/u", "Escritorios de trabajo en MDM/MDF con estructura metálica, cajoneras, para estaciones de trabajo de programación y administración."),
        ("3", "Sillas ergonómicas de oficina", "5", "Unidades", "—", "Sillas giratorias ergonómicas con soporte lumbar, reposabrazos ajustables y base de cinco ruedas."),
        ("4", "Mesa de juntas/conferencias", "1", "Unidad", "180 x 90 cm", "Mesa rectangular de reuniones con capacidad para 6 personas, superficie laminada."),
        ("5", "Estantería modular de almacenamiento", "2", "Unidades", "—", "Estantes modulares de 5 niveles cada uno, para almacenamiento de documentación, equipos y suministros de oficina."),
        ("6", "Archivador metálico de seguridad", "1", "Unidad", "4 gavetas", "Archivador vertical metálico con cerradura de seguridad, para custodia de documentos legales, financieros y contratos."),
        ("7", "Cortinas blackout para oficina", "3", "Unidades", "Según ventana", "Cortinas opacas reguladoras de luz para ventanas de la oficina, necesarias para producciones audiovisuales."),
    ]
    
    add_inventory_table(doc, oficina_items)
    
    doc.add_paragraph()
    
    # ---- TABLA 2: EQUIPOS DE CÓMPUTO Y TELECOMUNICACIONES ----
    add_centered(doc, "CATEGORÍA 2: EQUIPOS DE CÓMPUTO Y TELECOMUNICACIONES", bold=True, size=12)
    doc.add_paragraph()
    
    add_justified(doc,
        "Equipos de cómputo, telecomunicaciones, climatización e infraestructura "
        "tecnológica necesarios para el desarrollo de software, operaciones de "
        "plataforma y comunicaciones institucionales."
    )
    doc.add_paragraph()
    
    computo_items = [
        ("8", "Computadoras de escritorio / portátiles", "3", "Unidades", "—", "Equipos de cómputo para desarrollo de software, gestión administrativa y operaciones de plataforma. Incluyen procesador, memoria RAM y almacenamiento SSD."),
        ("9", "Monitores LED de alta resolución", "3", "Unidades", "24 pulgadas c/u", "Monitores de 24 pulgadas con resolución Full HD, panel IPS, para estaciones de trabajo de programación y diseño."),
        ("10", "Televisores LED inteligentes", "3", "Unidades", "55 pulgadas c/u", "Televisores Smart TV de 55 pulgadas con conexión WiFi, para monitoreo de dashboards, presentaciones de clientes y panel de operaciones en tiempo real."),
        ("11", "Teléfono institucional de línea", "1", "Unidad", "—", "Teléfono de línea fija con funciones de marcación rápida, contestadora y conferencia, para atención telefónica institucional."),
        ("12", "Equipos de aire acondicionado", "2", "Unidades", "12.000 BTU c/u", "Aires acondicionados split de 12.000 BTU, para climatización de la oficina principal y el área de producción audiovisual."),
        ("13", "Unidad de Protección Eléctrica (UPS)", "2", "Unidades", "1.500 VA c/u", "Sistemas de respaldo eléctrico continuo para protección de equipos de cómputo y producción ante fluctuaciones y cortes eléctricos."),
        ("14", "Router empresarial WiFi 6", "1", "Unidad", "—", "Router de alta capacidad con tecnología WiFi 6, soporte multiusuario y gestión de ancho de banda para operaciones simultáneas de desarrollo, producción y comunicaciones."),
        ("15", "Switch de red administrable", "1", "Unidad", "8 puertos", "Switch gigabit con 8 puertos Ethernet para interconexión de equipos de red en la oficina."),
        ("16", "Servidor NAS de almacenamiento", "1", "Unidad", "4 TB", "Dispositivo de almacenamiento en red con 4 TB de capacidad, para respaldos locales, archivos compartidos y gestión de contenido multimedia."),
        ("17", "Teléfonos celulares corporativos", "2", "Unidades", "—", "Teléfonos inteligentes asignados al equipo operativo para gestión de reservas, atención a clientes y notificaciones de plataforma."),
    ]
    
    add_inventory_table(doc, computo_items)
    
    doc.add_paragraph()
    
    # ---- TABLA 3: EQUIPOS DE AUDIOVISUAL Y PRODUCCIÓN ----
    add_centered(doc, "CATEGORÍA 3: EQUIPOS DE AUDIOVISUAL Y PRODUCCIÓN", bold=True, size=12)
    doc.add_paragraph()
    
    add_justified(doc,
        "Equipos destinados a la producción de contenido fotográfico y audiovisual "
        "para publicación de propiedades en la plataforma, material publicitario, "
        "marketing digital y contenido de redes sociales."
    )
    doc.add_paragraph()
    
    audiovisual_items = [
        ("18", "Cámara profesional Sony a6700", "1", "Unidad", "—", "Cámara mirrorless APS-C de 26 MP, sensor Exmor R CMOS, grabación de video 4K a 60fps, pantalla táctil articulada, conectividad WiFi/Bluetooth, para fotografía profesional de propiedades y contenido audiovisual."),
        ("19", "Lente Sigma Montura E 35mm f/1.4", "1", "Unidad", "35mm f/1.4", "Lente prime de distancia focal fija de 35mm y apertura máxima f/1.4, montura Sony E, ideal para fotografía de interiores, detalles arquitectónicos y retratos corporativos con desenfoque de fondo."),
        ("20", "Lente Sony 18-135mm Montura E", "1", "Unidad", "18-135mm f/3.5-5.6", "Lente zoom versátil de 18 a 135mm, apertura f/3.5-5.6, montura Sony E, ideal para fotografía general de propiedades, exteriores y tomas de gran angular a teleobjetivo."),
        ("21", "Estabilizador gimbal DJI RS 4 Pro", "1", "Unidad", "—", "Gimbal estabilizador de ejes para cámaras mirrorless, capacidad de carga hasta 4.5 kg, modo seguimiento, timelapse, vertical y sujetador integrado, para video fluido de recorridos virtuales por propiedades."),
        ("22", "Focos LED RGB profesionales", "2", "Unidades", "—", "Luces LED RGB con ajuste de temperatura de color (3.200K-5.600K), intensidad variable, control remoto por app, montura de estudio, para iluminación de escenarios de fotografía y video de interiores."),
        ("23", "Micrófonos profesionales Shure", "3", "Unidades", "—", "Micrófonos dinámicos cardioides profesionales marca Shure, con captación direccional de audio, para grabación de voz, entrevistas, contenido de redes sociales y podcasts corporativos."),
        ("24", "Trípode profesional de video", "1", "Unidad", "—", "Trípode de aluminio con fluid head, altura máxima 180 cm, capacidad de carga hasta 5 kg, bolsa de transporte incluida, para estabilización de tomas fijas en producciones de video."),
        ("25", "Pantalla de croma / green screen", "1", "Unidad", "2 x 3 metros", "Pantalla plegable de tela verde (chroma key) de 2 x 3 metros con soporte telescópico, para producción de contenido con fondos virtuales."),
        ("26", "Backdrops de estudio para fotografía", "2", "Unidades", "—", "Fondos de estudio de papel neutro (blanco y gris) con soporte extensible, para sesiones de fotografía profesional de productos del Boogie Store y material corporativo."),
        ("27", "Accesorios de iluminación y producción", "—", "Lote", "—", "Incluye: reflectores plegables, difusores, filtros ND, baterías recargables para cámara y gimbal, tarjetas de memoria SD de alta velocidad, cables de audio, maletines de transporte para equipos, limpiador de lentes y fundas protectoras."),
    ]
    
    add_inventory_table(doc, audiovisual_items)
    
    doc.add_paragraph()
    
    # ---- TABLA 4: ACTIVOS DIGITALES Y SOFTWARE ----
    add_centered(doc, "CATEGORÍA 4: ACTIVOS DIGITALES Y SOFTWARE", bold=True, size=12)
    doc.add_paragraph()
    
    add_justified(doc,
        "Activos intangibles de propiedad de la Sociedad, incluyendo la plataforma "
        "tecnológica, aplicaciones, dominios, sistemas de información y software "
        "desarrollado internamente o adquirido con licencia."
    )
    doc.add_paragraph()
    
    digital_items = [
        ("28", "Página web y plataforma principal Boogie", "1", "Sistema", "boogierent.com", "Plataforma web completa de intermediación de alojamiento vacacional, incluyendo: catálogo de propiedades con búsqueda avanzada, sistema de reservas, panel de anfitriones, panel de administración, chat en tiempo real, pasarela de pagos, sistema de reseñas, gestión de cupones, Boogie Store, sistema de ofertas 'Negocia tu Boogie', dashboard de métricas y estadísticas."),
        ("29", "Aplicaciones móviles Boogie (iOS y Android)", "2", "Aplicaciones", "—", "Aplicaciones nativas multiplataforma (iOS y Android) con funcionalidades de búsqueda, reserva, check-in/check-out, notificaciones push, chat, pagos y gestión de propiedades."),
        ("30", "Backend API en Go microservice", "1", "Sistema", "—", "Servicio backend en lenguaje Go con arquitectura de microservicios, base de datos PostgreSQL (Supabase), autenticación JWT, procesamiento de pagos cripto (CryptAPI), integración BCV para tipos de cambio y webhooks de verificación (MetaMap)."),
        ("31", "Base de datos PostgreSQL (Supabase)", "1", "Sistema", "—", "Base de datos relacional alojada en Supabase con tablas de: usuarios, propiedades, reservas, pagos, reseñas, cupones, ofertas, notificaciones, audit logs y configuraciones de plataforma."),
        ("32", "Dominio boogierent.com", "1", "Unidad", "—", "Nombre de dominio principal registrado y administrado por la Sociedad para la operación de la plataforma web."),
        ("33", "Certificado SSL/TLS", "1", "Unidad", "—", "Certificado de seguridad digital para encriptación de comunicaciones entre usuarios y la plataforma."),
        ("34", "Marcas registradas y logotipos", "1", "Sistema", "—", "Incluye: logotipo principal Boogie, logotipo Boogie Store, logotipo Negocia tu Boogie, tipografías institucionales, paleta de colores, guía de marca y activos gráficos para redes sociales."),
        ("35", "Licencias de software y servicios SaaS", "—", "Lote", "—", "Incluye: Vercel (hosting), Supabase (base de datos), Google Cloud/AWS, Mailgun/SendGrid (correos), Google Analytics, plan Ultra de Boogie, MetaMap (verificación), CryptAPI (pagos crypto) y demás herramientas de productividad y operación."),
    ]
    
    add_inventory_table(doc, digital_items)
    
    doc.add_paragraph()
    
    # ---- TABLA 5: EQUIPOS DE OFICINA COMPLEMENTARIOS ----
    add_centered(doc, "CATEGORÍA 5: EQUIPOS Y SUMINISTROS DE OFICINA", bold=True, size=12)
    doc.add_paragraph()
    
    add_justified(doc,
        "Equipos complementarios, suministros y elementos de seguridad industrial "
        "y protección civil necesarios para el cumplimiento de las normativas "
        "laborales vigentes (LOPCYMAT) y el adecuado funcionamiento de la sede "
        "operativa de la Sociedad."
    )
    doc.add_paragraph()
    
    oficina_comp_items = [
        ("36", "Impresora multifuncional láser", "1", "Unidad", "—", "Impresora láser multifuncional (impresión, copia, escáner, fax), resolución 1200 dpi, conectividad WiFi y USB, para documentación legal, contratos y facturación."),
        ("37", "Escáner de documentos profesional", "1", "Unidad", "—", "Escáner de alimentación automática de documentos (ADF) con velocidad de 30 ppm, para digitalización de cédulas, comprobantes bancarios y contratos firmados."),
        ("38", "Destructora de documentos", "1", "Unidad", "—", "Destructora de papel con seguridad nivel P-4 (corte cruzado), capacidad de 10 hojas por pasada, para destrucción segura de documentos confidenciales y datos sensibles."),
        ("39", "Dispensador de agua", "1", "Unidad", "—", "Dispensador de agua con botellón recargable, con opciones de agua fría y temperatura ambiente, para uso del equipo operativo."),
        ("40", "Kit de suministros de oficina", "1", "Lote", "—", "Incluye: resmas de papel, bolígrafos, marcadores, clips, grapadoras, perforadoras, cinta adhesiva, carpetas archivadoras, pizarras blancas, borradores, notas adhesivas y demas material fungible de oficina."),
        ("41", "Botiquín de primeros auxilios", "1", "Unidad", "—", "Kit de primeros auxilios con insumos básicos de emergencia (vendajes, antisépticos, guantes), para cumplimiento de normativa de seguridad laboral (LOPCYMAT)."),
        ("42", "Extintor de incendios", "1", "Unidad", "—", "Extintor multipropósito ABC de 4.5 kg, con soporte de pared y registro de inspección, para cumplimiento de normativa de protección civil."),
    ]
    
    add_inventory_table(doc, oficina_comp_items)
    
    doc.add_paragraph()
    
    # Declaración del inventario
    add_justified(doc,
        "Los comparecientes declaran, bajo juramento, que los activos descritos en "
        "el presente inventario pertenecen legítimamente a los fundadores de la Sociedad, "
        "se encuentran en buen estado de funcionamiento y son suficientes para el inicio "
        "de las actividades contempladas en el objeto social. Dichos activos serán "
        "transferidos a la Sociedad una vez esta quede legalmente constituida, quedando "
        "registrados en el libro de inventario conforme a lo dispuesto en el artículo 329 "
        "del Código de Comercio."
    )
    
    doc.add_paragraph()
    
    # ==================== ADMINISTRACION ====================
    add_centered(doc, "ADMINISTRACIÓN DE LA SOCIEDAD", bold=True, size=14)
    doc.add_paragraph()
    
    add_article_multi(doc, "CLÁUSULA OCTAVA:", [
        "La administración de la Sociedad estará a cargo de una Junta Directiva compuesta "
        "por CUATRO (4) miembros, quienes serán nombrados por la Asamblea de Accionistas "
        "y ejercerán sus funciones por tiempo indefinido, mientras la Asamblea no disponga "
        "su remoción o sustitución.",
        
        "La Junta Directiva quedará integrada de la siguiente manera:"
    ])
    
    # Tabla de Junta Directiva
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    hdr = table.rows[0]
    hdr.cells[0].text = "CARGO"
    hdr.cells[1].text = "NOMBRE COMPLETO / CÉDULA"
    for cell in hdr.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = FONT_NAME
                run.font.size = Pt(11)
    
    # Data
    data = [
        ("Presidente", "Sebastián Chacón Peña\nV-27.333.855"),
        ("Vicepresidente", "Nuryvel Antonieta Peña Gonzalez\nV-10.505.412"),
        ("Director / CTO", "Javier Viloria Ferrer\nV-28.567.944"),
        ("Secretaria", "Yennuri Peña Gonzalez\nV-10.678.456"),
    ]
    
    for i, (cargo, nombre) in enumerate(data):
        row = table.rows[i + 1]
        row.cells[0].text = cargo
        row.cells[1].text = nombre
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA NOVENA:",
        "La Junta Directiva se reunirá cuando sea convocada por el Presidente o por "
        "cualquier Director, previo aviso con al menos 48 horas de anticipación. El quórum "
        "para sesionar será de la mitad más uno de sus miembros, y las decisiones se tomarán "
        "por mayoría simple de los presentes. En caso de empate, el voto del Presidente será "
        "dirimente."
    )
    
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA DÉCIMA:",
        "La Junta Directiva tendrá las siguientes atribuciones entre otras: "
        "ejercer la administración de la Sociedad, representarla judicial y extrajudicialmente "
        "cuando medie delegación expresa del Presidente, ejecutar las resoluciones de la "
        "Asamblea de Accionistas, celebrar contratos, otorgar poderes especiales con "
        "facultades limitadas, aprobar presupuestos, y en general realizar todos los actos "
        "de administración necesarios para el cumplimiento del objeto social."
    )
    
    doc.add_paragraph()
    
    # ==================== REPRESENTACION LEGAL ====================
    add_centered(doc, "REPRESENTACIÓN LEGAL", bold=True, size=14)
    doc.add_paragraph()
    
    add_article_multi(doc, "CLÁUSULA DÉCIMA PRIMERA:", [
        "La representación legal de la Sociedad será ejercida por el Presidente y la "
        "Vicepresidenta de la Junta Directiva, quienes podrán actuar individualmente "
        "o conjuntamente conforme a las facultades que se establecen a continuación.",
        
        "El ciudadano Sebastián Chacón Peña, Presidente, titular de la cédula de "
        "identidad Nro. V-27.333.855, queda investido de poderes amplios para representar "
        "a la Sociedad ante cualquier persona natural o jurídica, ente público o privado, "
        "nacional o internacional, con las siguientes facultades:",
        
        "a) Representar a la Sociedad ante toda clase de autoridades judiciales, "
        "administrativas, fiscales, bancarias, municipales y en general ante cualquier "
        "persona natural o jurídica.",
        
        "b) Celebrar, modificar, rescindir y ejecutar todo tipo de contratos, convenios "
        "y acuerdos.",
        
        "c) Abrir, operar y cerrar cuentas bancarias corrientes, de ahorro, depósitos "
        "a plazo fijo y cualquier otro instrumento financiero.",
        
        "d) Realizar operaciones bancarias, incluyendo transferencias, pagos, "
        "depósitos, retiros y emisión de cheques.",
        
        "e) Suscribir y negociar todo tipo de títulos valores.",
        
        "f) Administrar los bienes de la Sociedad, adquirir y enajenar muebles e inmuebles.",
        
        "g) Contratar, administrar y despachar al personal de la Sociedad.",
        
        "h) Cumplir todas las obligaciones formales y sustantivas ante organismos tributarios "
        "nacionales, estadales y municipales, incluyendo SENIAT.",
        
        "i) Otorgar poderes especiales con facultades limitadas a terceros de confianza.",
        
        "j) Realizar cualquier acto de gestión necesario para el cumplimiento del objeto social."
    ])
    
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA DÉCIMA SEGUNDA:",
        "La ciudadana Nuryvel Antonieta Peña Gonzalez, Vicepresidenta, titular de la cédula "
        "de identidad Nro. V-10.505.412, queda igualmente investida de los mismos poderes "
        "amplios de representación legal contenidos en la cláusula anterior, pudiendo actuar "
        "individual o conjuntamente con el Presidente, con idénticas facultades."
    )
    
    doc.add_paragraph()
    
    # ==================== ASAMBLEAS ====================
    add_centered(doc, "ASAMBLEA DE ACCIONISTAS", bold=True, size=14)
    doc.add_paragraph()
    
    add_article_multi(doc, "CLÁUSULA DÉCIMA TERCERA:", [
        "La Asamblea de Accionistas es el organismo supremo de la Sociedad y será "
        "integrada por todos los accionistas legalmente reconocidos.",
        
        "Se distinguirán asambleas ordinarias y extraordinarias:"
    ])
    
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA DÉCIMA CUARTA:",
        "La Asamblea Ordinaria se reunirá por lo menos una vez al ano dentro de los "
        "primeros CUATRO (4) meses posteriores al cierre de cada ejercicio social, "
        "con el objeto de: examinar los informes y balances del ejercicio anterior, "
        "aprobar o desaprobar la gestión de la Junta Directiva y distribuir utilidades, "
        "así como nombrar y remover directores. El ejercicio social abarca del 1 de "
        "enero al 31 de diciembre de cada ano."
    )
    
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA DÉCIMA QUINTA:",
        "La Asamblea Extraordinaria se reunirá cuando la convoque la Junta Directiva "
        "o accionistas que representen por lo menos la quinta parte del capital suscrito, "
        "y tendrá competencia para: modificar los Estatutos Sociales, aumentar o reducir "
        "el capital, autorizar la emisión de bonos, decidir la disolución anticipada de "
        "la Sociedad, fusionarse con otra empresa, y en general resolver todos los asuntos "
        "de especial gravedad."
    )
    
    doc.add_paragraph()
    
    # ==================== UTILIDADES ====================
    add_centered(doc, "UTILIDADES Y PÉRDIDAS", bold=True, size=14)
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA DÉCIMA SEXTA:",
        "Las utilidades netas que resulten del balance general de cada ejercicio social "
        "serán distribuidas entre los accionistas en proporcion directa a sus respectivas "
        "participaciones accionarias, previa constitución de la reserva legal del diez por "
        "ciento (10%) exigida por el Código de Comercio, y las demas reservas que la "
        "Asamblea determine."
    )
    
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA DÉCIMA SÉPTIMA:",
        "En caso de pérdida, ésta se distribuirá entre los accionistas en proporcion "
        "directa a sus respectivas participaciones accionarias, salvo que la Asamblea "
        "resuelva cubrirla con las reservas existentes."
    )
    
    doc.add_paragraph()
    
    # ==================== AUDITORIA ====================
    add_centered(doc, "AUDITORÍA", bold=True, size=14)
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA DÉCIMA OCTAVA:",
        "La Asamblea de Accionistas podrá designar Comisarios o auditores externos para "
        "examinar los libros, cuentas y operaciones de la Sociedad. Los Comisarios "
        "tendrán las facultades conferidas por el Código de Comercio y presentarán su "
        "informe ante la Asamblea Ordinaria."
    )
    
    doc.add_paragraph()
    
    # ==================== DISOLUCION ====================
    add_centered(doc, "DISOLUCIÓN Y LIQUIDACIÓN", bold=True, size=14)
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA DÉCIMA NOVENA:",
        "La Sociedad se disolverá por las causales establecidas en el artículo 338 del "
        "Código de Comercio, en especial: por decisión de la Asamblea Extraordinaria "
        "con el voto favorable de accionistas que representen por lo menos las tres cuartas "
        "partes del capital social, por vencimiento del término de duración, por pérdida de "
        "las tres cuartas partes del capital social, o por cualquiera de las demas causales "
        "previstas en la Ley."
    )
    
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA VIGÉSIMA:",
        "En caso de disolución, la Sociedad entrará en liquidación. Los liquidadores serán "
        "nombrados por la Asamblea y se encargarán de cobrar los créditos, pagar las deudas "
        "y distribuir el remanente entre los accionistas en proporcion a sus participaciones "
        "accionarias. Las operaciones de liquidación no podrán exceder el plazo de DOS (2) "
        "anos contados desde la fecha de disolución."
    )
    
    doc.add_paragraph()
    
    # ==================== DISPOSICIONES GENERALES ====================
    add_centered(doc, "DISPOSICIONES GENERALES", bold=True, size=14)
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA VIGÉSIMA PRIMERA:",
        "Para todo lo no previsto en estos Estatutos, se aplicarán supletoriamente las "
        "disposiciones del Código de Comercio de la República Bolivariana de Venezuela y "
        "demás leyes aplicables."
    )
    
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA VIGÉSIMA SEGUNDA:",
        "Todas las notificaciones entre la Sociedad y sus accionistas se realizarán por "
        "escrito a las direcciones registradas o por medios electrónicos cuando las partes "
        "así lo acuerden."
    )
    
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA VIGÉSIMA TERCERA:",
        "La transferencia de acciones de la Sociedad será libre entre accionistas. "
        "Para transferencias a terceros, se requerirá la aprobación "
        "previa de la Junta Directiva y la notificación formal ante el Registro Mercantil."
    )
    
    doc.add_paragraph()
    
    # ==================== JURISDICCION ====================
    add_centered(doc, "JURISDICCIÓN", bold=True, size=14)
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA VIGÉSIMA CUARTA:",
        "Las controversias que se susciten entre los accionistas, entre accionistas y "
        "la Sociedad, o entre la Sociedad y terceros con motivo de los presentes "
        "Estatutos Sociales, serán resueltas por los Juzgados de Primera Instancia "
        "en lo Mercantil de la Circunscripción Judicial del Distrito Capital, o en "
        "su defecto, por los tribunales competentes de Caracas, con renuncia expresa "
        "a cualquier otro fuero que pudiera corresponder."
    )
    
    doc.add_paragraph()
    
    # ==================== INSCRIPCION REGISTRAL ====================
    add_centered(doc, "INSCRIPCIÓN EN EL REGISTRO MERCANTIL", bold=True, size=14)
    doc.add_paragraph()
    
    add_article(doc, "CLÁUSULA VIGÉSIMA QUINTA:",
        "Los presentes Estatutos Sociales serán inscritos en el Registro Mercantil del "
        "Distrito Capital o del lugar de domicilio de la Sociedad, dentro de los QUINCE "
        "(15) días hábiles siguientes a su otorgamiento, conforme a lo establecido en el "
        "artículo 202 del Código de Comercio."
    )
    
    doc.add_paragraph()
    
    # ==================== SOTSCRICION Y PAGO ====================
    add_centered(doc, "CONSTANCIA DE SUSCRIPCIÓN Y PAGO DEL CAPITAL", bold=True, size=14)
    doc.add_paragraph()
    
    add_justified(doc,
        "Los comparecientes declaran, de forma conjunta y solidaria, bajo juramento "
        "de decir la verdad y a los fines del artículo 200 del Código de Comercio, "
        "que:"
    )
    
    doc.add_paragraph()
    
    add_justified(doc,
        "PRIMERO: Que el capital de la Sociedad ha sido íntegramente suscrito por "
        "los accionistas fundadores en las proporciones indicadas en la CLÁUSULA QUINTA "
        "de estos Estatutos."
    )
    
    doc.add_paragraph()
    
    add_justified(doc,
        "SEGUNDO: Que el capital suscrito ha sido íntegramente pagado por los "
        "accionistas fundadores al momento de la celebración del presente acto constitutivo, "
        "en dólares de los Estados Unidos de América (USD), o su equivalente en bolívares "
        "según la tasa oficial del Banco Central de Venezuela, conforme a los comprobantes "
        "de pago correspondientes."
    )
    
    doc.add_paragraph()
    
    add_justified(doc,
        "TERCERO: Que no se ha hecho uso de la facultad prevista en el artículo 204 "
        "del Código de Comercio relativa a la suscripción sucesiva del capital social."
    )
    
    doc.add_paragraph()
    
    # ==================== DECLARACION DE VOLUNTAD FINAL ====================
    add_justified(doc,
        "Los comparecientes, conscientes del contenido y alcance legal de la presente "
        "acta, declaran su firme y expresa voluntad de constituir la Sociedad denominada "
        "BOOGIE RENT C.A., bajo los términos y condiciones contenidos en este documento, "
        "sometiéndose a las disposiciones legales de la República Bolivariana de Venezuela "
        "y a la jurisdicción de los tribunales de la ciudad de Caracas, Distrito Capital."
    )
    
    doc.add_paragraph()
    
    add_justified(doc,
        "Se eleva a escritura pública el presente instrumento y se ordena su inserción "
        "en el libro de correspondencia del Registrador Público que suscribe."
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ==================== FIRMAS ====================
    add_centered(doc, "SUSCRITAS Y FIRMADAS", bold=True, size=14)
    add_centered(doc, "POR LOS COMPARECIENTES", bold=False, size=12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Firma Socio 1
    add_signature_line(doc, "Sebastián Chacón Peña", "27.333.855", 
                       "Socio Fundador — Presidente")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Firma Socio 2
    add_signature_line(doc, "Nuryvel Antonieta Peña Gonzalez", "10.505.412", 
                       "Socia Fundadora — Vicepresidenta")
    
    doc.add_paragraph()
    
    # ==================== DIRECTORES DESIGNADOS ====================
    add_centered(doc, "DIRECTORES DESIGNADOS", bold=True, size=14)
    add_centered(doc, "(Aceptan sus cargos)", bold=False, size=12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Director / CTO
    add_signature_line(doc, "Javier Viloria Ferrer", "28.567.944", 
                       "Director / CTO")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Secretaria
    add_signature_line(doc, "Yennuri Peña Gonzalez", "10.678.456", 
                       "Secretaria")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ==================== NOTARIO ====================
    add_centered(doc, "AUTORIZA", bold=True, size=14)
    add_centered(doc, "EL REGISTRADOR PÚBLICO", bold=False, size=12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_________________________________")
    run.font.name = FONT_NAME
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("[Nombre del Registrador Público]")
    run2.font.name = FONT_NAME
    run2.italic = True
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("[Número de Registro del Registrador]")
    run3.font.name = FONT_NAME
    run3.font.size = Pt(11)
    run3.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ==================== NOTA DE INSCRIPCION ====================
    add_centered(doc, "INSCRIPCIÓN EN EL REGISTRO MERCANTIL", bold=True, size=12)
    
    doc.add_paragraph()
    
    add_justified(doc,
        "Vistos los presentes Estatutos, se ordena la inscripción en el Registro Mercantil "
        "del Distrito Capital / Circunscripción Judicial correspondiente, de conformidad con "
        "lo establecido en el artículo 202 del Código de Comercio de Venezuela."
    )
    
    doc.add_paragraph()
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run("Fecha de inscripción: ____________________")
    run_date.font.name = FONT_NAME
    run_date.italic = True
    
    doc.add_paragraph()
    
    p_reg = doc.add_paragraph()
    p_reg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_reg = p_reg.add_run("Tomo: _________ Número: _________ Folio: _________")
    run_reg.font.name = FONT_NAME
    run_reg.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ==================== FOOTER NOTE ====================
    add_centered(doc, "— DOCUMENTO GENERADO —", bold=True, size=10)
    add_centered(doc, "Este documento es una plantilla base para el registro mercantil.", size=9)
    add_centered(doc, "Requiere revisión y firma ante Registrador Público en Venezuela.", size=9)
    add_centered(doc, "Contacto legal: legal@boogie.com.ve", size=9)
    
    # Save
    doc.save(OUTPUT_FILE)
    print(f"Documento generado exitosamente: {OUTPUT_FILE}")
    print(f"Tamaño: {os.path.getsize(OUTPUT_FILE):,} bytes")


if __name__ == "__main__":
    generate()
