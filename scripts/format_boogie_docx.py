import sys
import io
import copy
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

INPUT = r'C:\Users\swmfi\Documents\Proyectos\swmTech\boogie\docs\legal\BOOGIE.docx'
OUTPUT = r'C:\Users\swmfi\Documents\Proyectos\swmTech\boogie\docs\legal\BOOGIE_Formateado.docx'

BRAND_PRIMARY = RGBColor(0x1B, 0x43, 0x32)
BRAND_ACCENT = RGBColor(0xE7, 0x6F, 0x51)
BRAND_GREEN_LIGHT = RGBColor(0x52, 0xB7, 0x88)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
MID_TEXT = RGBColor(0x4B, 0x45, 0x40)
LIGHT_TEXT = RGBColor(0x6B, 0x65, 0x60)

FONT_BODY = 'Calibri'
FONT_HEADING = 'Calibri'


def set_paragraph_spacing(para, before=0, after=0, line_spacing=1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_horizontal_line(doc, color='1B4332'):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    set_paragraph_spacing(p, before=4, after=4)
    return p


def add_styled_paragraph(doc, text, font_size=11, bold=False, italic=False,
                          color=DARK_TEXT, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          font_name=FONT_BODY, before=0, after=6, line_spacing=1.15,
                          space_after=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = font_name
    set_paragraph_spacing(p, before=before, after=space_after if space_after is not None else after,
                          line_spacing=line_spacing)
    return p


def add_section_title(doc, text):
    add_horizontal_line(doc, '1B4332')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BRAND_PRIMARY
    run.font.name = FONT_HEADING
    set_paragraph_spacing(p, before=12, after=4, line_spacing=1.2)
    add_horizontal_line(doc, '1B4332')
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=2, after=2)


def add_subsection_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = BRAND_PRIMARY
    run.font.name = FONT_HEADING
    set_paragraph_spacing(p, before=16, after=6, line_spacing=1.2)
    return p


def add_sub_subsection(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F)
    run.font.name = FONT_HEADING
    set_paragraph_spacing(p, before=10, after=4, line_spacing=1.15)
    return p


def add_body(doc, text, before=2, after=6, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_TEXT
    run.font.name = FONT_BODY
    set_paragraph_spacing(p, before=before, after=after, line_spacing=1.15)
    return p


def add_bullet(doc, text, indent_level=0, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    indent = Inches(0.3 + indent_level * 0.3)
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = Inches(-0.2)
    bullet_run = p.add_run('•  ')
    bullet_run.font.size = Pt(10.5)
    bullet_run.font.color.rgb = BRAND_GREEN_LIGHT
    bullet_run.font.name = FONT_BODY
    if bold_prefix:
        bp = p.add_run(bold_prefix)
        bp.font.size = Pt(10.5)
        bp.font.bold = True
        bp.font.color.rgb = DARK_TEXT
        bp.font.name = FONT_BODY
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_TEXT
    run.font.name = FONT_BODY
    set_paragraph_spacing(p, before=1, after=3, line_spacing=1.1)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = FONT_BODY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B4332"/>')
        cell._element.get_or_add_tcPr().append(shading)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = FONT_BODY
            run.font.color.rgb = DARK_TEXT
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bg = 'FFFFFF' if r_idx % 2 == 0 else 'F0FDF4'
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
            cell._element.get_or_add_tcPr().append(shading)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="D8F3DC"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="D8F3DC"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D8F3DC"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="D8F3DC"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D8F3DC"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="D8F3DC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    p_after = doc.add_paragraph()
    set_paragraph_spacing(p_after, before=2, after=8)
    return table


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.8)
    p.paragraph_format.right_indent = Inches(0.8)
    run = p.add_run(f'«{text}»')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.bold = True
    run.font.color.rgb = BRAND_PRIMARY
    run.font.name = FONT_BODY
    set_paragraph_spacing(p, before=12, after=12, line_spacing=1.3)
    return p


def add_spacer(doc, before=6):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=before, after=0)


doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = FONT_BODY
style.font.size = Pt(10.5)
style.font.color.rgb = DARK_TEXT
style.paragraph_format.line_spacing = 1.15

# ═══════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════
for _ in range(4):
    add_spacer(doc, 12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BOOGIE')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = BRAND_PRIMARY
run.font.name = FONT_HEADING
set_paragraph_spacing(p, before=0, after=2)

add_horizontal_line(doc, '52B788')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PLAN DE NEGOCIO E INVERSIÓN')
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2D, 0x6A, 0x4F)
run.font.name = FONT_HEADING
set_paragraph_spacing(p, before=8, after=24)

for _ in range(3):
    add_spacer(doc, 8)

cover_lines = [
    ('Documento Confidencial', 12, True, BRAND_PRIMARY),
    ('Versión 1.0 — 2025', 11, False, MID_TEXT),
    ('Preparado para Entidades Financieras', 11, False, MID_TEXT),
    ('Clasificación: Confidencial — Uso Exclusivo del Destinatario', 9.5, True, BRAND_ACCENT),
]
for text, size, bold, color in cover_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = FONT_BODY
    set_paragraph_spacing(p, before=2, after=2)

doc.add_page_break()

# ═══════════════════════════════════════════════
# ÍNDICE
# ═══════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ÍNDICE')
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = BRAND_PRIMARY
run.font.name = FONT_HEADING
set_paragraph_spacing(p, before=8, after=12)

add_horizontal_line(doc, '52B788')
add_spacer(doc, 4)

toc_items = [
    ('01.', 'Resumen Ejecutivo', 'Visión general del proyecto, tesis de inversión y métricas principales.'),
    ('02.', 'Descripción de la Plataforma', 'Funcionalidades técnicas, diferenciadores y ventajas competitivas.'),
    ('03.', 'Modelos de Negocio', 'Cinco fuentes de ingresos: comisiones, suscripciones, commodities, servicios, flipping.'),
    ('04.', 'Flipping Inmobiliario y Estrategia de Activos', 'Modelo principal, proyecciones del piloto y estrategia de cartera.'),
    ('05.', 'Ruta Estratégica', 'Hoja de ruta a 24 meses con hitos operativos y financieros.'),
    ('06.', 'Proyecciones Financieras', 'Proyecciones por línea de negocio, métricas de crecimiento y break-even.'),
    ('07.', 'Objetivos del Financiamiento', 'Asignación de capital por escenarios y retorno esperado.'),
    ('08.', 'Propuesta al Socio Bancario', 'Beneficios, estructura de relación y condiciones propuestas.'),
]

for num, title, desc in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.3)
    run_num = p.add_run(f'{num}  ')
    run_num.font.size = Pt(11)
    run_num.font.bold = True
    run_num.font.color.rgb = BRAND_PRIMARY
    run_num.font.name = FONT_BODY
    run_title = p.add_run(title)
    run_title.font.size = Pt(11)
    run_title.font.bold = True
    run_title.font.color.rgb = DARK_TEXT
    run_title.font.name = FONT_BODY
    set_paragraph_spacing(p, before=6, after=0, line_spacing=1.2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.left_indent = Inches(0.6)
    run_d = p2.add_run(desc)
    run_d.font.size = Pt(9.5)
    run_d.font.italic = True
    run_d.font.color.rgb = MID_TEXT
    run_d.font.name = FONT_BODY
    set_paragraph_spacing(p2, before=0, after=8, line_spacing=1.1)

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECCIÓN 01 — RESUMEN EJECUTIVO
# ═══════════════════════════════════════════════
add_section_title(doc, 'SECCIÓN 01 — RESUMEN EJECUTIVO')

add_body(doc, 'Boogie es una plataforma de hospitalidad de corta estancia que opera como ecosistema integral en el mercado venezolano. Su diseño estratégico la posiciona no únicamente como un marketplace de alquileres vacacionales, sino como el motor de entrada hacia dos modelos de negocio de alto retorno: la inversión y flipping inmobiliario, y la venta al detal de bienes de consumo.')

add_body(doc, 'La plataforma constituye el medio de generación de confianza y tráfico de mercado que alimenta dos modelos de negocio principales: la inversión y flipping inmobiliario, y la venta al detal de bienes de consumo (commodities). Si bien la plataforma per se significará una entrada de ingresos significativa en las arcas de la empresa, las utilidades de la misma tendrán el fin único y último de alimentar los dos modelos de negocios principales, que garantizarán la solvencia económica de la empresa ante cualquier circunstancia, al contar con activos tangibles bien posicionados y capaces de generar dividendos fijos.')

add_subsection_title(doc, 'Métricas Principales')

add_bullet(doc, '', bold_prefix='Fuentes de ingreso: 5')
add_bullet(doc, '', bold_prefix='ROI mensual estimado: 18%')
add_bullet(doc, '', bold_prefix='Rango de financiamiento: $35,000 – $55,000')
add_bullet(doc, '', bold_prefix='Break-even proyectado: 14 – 18 meses')

add_spacer(doc, 6)

add_body(doc, 'Tesis de inversión: El financiamiento solicitado se destina a tres objetivos concretos: adquisición de un inmueble piloto, despliegue de máquinas expendedoras en dos residenciales aliados, y ejecución de una campaña de adquisición de usuarios. El 55%–60% del capital se convierte directamente en un activo inmobiliario con valor de mercado verificable.')

add_quote(doc, 'La plataforma genera confianza. La confianza genera tráfico. El tráfico genera capital. El capital adquiere activos. Los activos garantizan solvencia.')

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECCIÓN 02 — DESCRIPCIÓN DE LA PLATAFORMA
# ═══════════════════════════════════════════════
add_section_title(doc, 'SECCIÓN 02 — DESCRIPCIÓN DE LA PLATAFORMA')

add_body(doc, 'Boogie es la primera plataforma de alquiler de corta estancia en Venezuela que integra procesamiento de pagos en bolívares, criptomonedas y divisas dentro de un mismo ecosistema, combinando gestión de propiedades, servicios de atención personalizada y una tienda integrada de productos y servicios bajo demanda.')

add_subsection_title(doc, 'Características Principales')

add_sub_subsection(doc, 'Procesamiento Multimoneda')
add_body(doc, 'Procesamiento nativo en bolívares, USDT, BTC, ETH y divisas fiat. Sin conversiones manuales ni intermediarios externos para el usuario final. Métodos disponibles: Pago Móvil, transferencia bancaria, Zelle, USDT TRC20 (vía CryptAPI con confirmación automática en blockchain), efectivo en puntos autorizados y Boogie Wallet.')

add_sub_subsection(doc, 'Gestión del Anfitrión')
add_body(doc, 'Dashboard con métricas de ocupación, pricing dinámico, gestión de calendario, herramientas de posicionamiento y reportes de rendimiento. Sistema de verificación de identidad (MetaMap), gestión de reservas con estados avanzados, chat integrado con huéspedes y sistema de reseñas bidireccional.')

add_sub_subsection(doc, 'Experiencia del Huésped')
add_body(doc, 'Check-in digital, concierge personalizado, servicios bajo demanda integrados (consolas, chefs, maquilladores, peluqueros), sistema de recomendaciones y soporte en tiempo real. Flujo de reserva en 4 pasos: resumen, "Arma tu Boogie" (productos y servicios), pago y confirmación.')

add_sub_subsection(doc, 'Infraestructura Técnica')
add_body(doc, 'Arquitectura dual: frontend Next.js 16 con App Router + backend Go con API REST. Base de datos PostgreSQL en Supabase con 20+ entidades y 60+ endpoints. Autenticación JWT con Supabase Auth, verificación JWKS, rate limiting, CORS configurado y pipeline de CI/CD con Husky + lint-staged. Desplegado en Vercel (frontend) y servicio dedicado (backend).')

add_subsection_title(doc, 'Ventajas Competitivas')

add_sub_subsection(doc, 'Tecnológicas')
add_bullet(doc, 'Procesamiento de pagos en bolívares — funcionalidad no disponible en competidores')
add_bullet(doc, 'Integración cripto nativa (USDT, BTC, ETH)')
add_bullet(doc, 'Arquitectura mobile-first con experiencia de usuario optimizada')
add_bullet(doc, 'API abierta para integraciones con sistemas de terceros')
add_bullet(doc, 'Algoritmo de pricing dinámico basado en demanda local')
add_bullet(doc, 'Inteligencia de mercado con datos en tiempo real')

add_sub_subsection(doc, 'Operativas')
add_bullet(doc, 'Servicio de concierge y atención personalizada al huésped')
add_bullet(doc, 'Red integrada de servicios complementarios')
add_bullet(doc, 'Modelo de suscripción escalable para anfitriones')
add_bullet(doc, 'Equipo de producción audiovisual interno')
add_bullet(doc, 'Estrategia vertical: plataforma, commodities e inmuebles')
add_bullet(doc, 'Integración con máquinas expendedoras propias')

add_spacer(doc, 6)
add_body(doc, 'Ningún competidor en el mercado venezolano combina procesamiento en bolívares, servicios de comodidad integrados, venta de commodities y estrategia de inversión inmobiliaria dentro de un mismo ecosistema.')

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECCIÓN 03 — MODELOS DE NEGOCIO
# ═══════════════════════════════════════════════
add_section_title(doc, 'SECCIÓN 03 — MODELOS DE NEGOCIO')

add_body(doc, 'Boogie opera bajo cinco fuentes de ingreso complementarias, diseñadas para generar flujos de caja diversificados y reducir la dependencia de cualquier línea individual.')

# 2.1
add_horizontal_line(doc, '52B788')
add_subsection_title(doc, 'Fuente 2.1 — Comisiones por Intermediación')

add_body(doc, 'Boogie cobra una comisión del 12% al 18% sobre cada transacción completada, actuando como intermediario en el alquiler de corta estancia de propiedades de terceros y en la reserva de canchas deportivas. Esta comisión se divide entre huésped (6%) y anfitrión (3%), con la diferencia configurada dinámicamente según el tipo de propiedad y demanda.')

add_bullet(doc, '', bold_prefix='Comisión promedio: 15%')
add_bullet(doc, '', bold_prefix='Ingreso mensual est. mes 6: $2,800')
add_bullet(doc, '', bold_prefix='Margen bruto: 85%')

# 2.2
add_horizontal_line(doc, '52B788')
add_subsection_title(doc, 'Fuente 2.2 — Suscripciones Premium para Anfitriones')

add_body(doc, 'Boogie ofrece niveles de suscripción mensual que otorgan a los anfitriones posicionamiento privilegiado, atención personalizada, fotografía profesional y acceso a herramientas avanzadas de gestión. Este modelo genera ingreso recurrente predecible, independiente de la tasa de ocupación.')

add_table(doc,
    ['Plan', 'Mensualidad', 'Alcance', 'Proyección Año 1'],
    [
        ['Básico', '$15/mes', 'Listado estándar, analytics básico', '$1,800/mes'],
        ['Pro', '$45/mes', 'Top 3 en resultados, concierge, fotografía', '$2,700/mes'],
        ['Elite', '$99/mes', 'Todo Pro + pricing dinámico + manager', '$1,980/mes'],
        ['TOTAL', '', '', '$6,480/mes'],
    ]
)

add_body(doc, 'Con una tasa de retención mensual del 82%, el ingreso recurrente anual proyectado (ARR) asciende a $77,760, con un tiempo de vida útil del cliente (LTV) superior a los $380 por anfitrión. (120 anfitriones activos al cierre del Año 1.)')

add_body(doc, 'Importancia estratégica: Las suscripciones constituyen un flujo de ingresos predecible, independiente de la tasa de ocupación. Representan el flujo de caja operativo que sostiene la actividad mientras los modelos principales maduran.')

# 2.3
add_horizontal_line(doc, '52B788')
add_subsection_title(doc, 'Fuente 2.3 — Venta de Bienes de Consumo')

add_body(doc, '"Arma Tu Boogie" es un servicio de personalización pre-llegada que permite al huésped seleccionar y pre-pagar un conjunto de productos que será dispuesto en la propiedad antes del check-in. El huésped accede a un catálogo de snacks, bebidas, artículos de higiene, cuidado personal y consumibles, organizado por categorías con precios en USD y Bolívares. Este servicio se integra directamente en el flujo de reserva como segundo paso, garantizando visibilidad y alta tasa de conversión.')

add_body(doc, 'Este canal se complementa con máquinas expendedoras de última generación — operadas bajo la marca y estándares de calidad de Boogie — instaladas en edificios residenciales aliados. Las máquinas operan con pantalla táctil, múltiples métodos de pago y reposición programada bajo el label de Boogie.')

add_sub_subsection(doc, 'Servicio "Arma Tu Boogie"')
add_bullet(doc, 'Configuración del pedido previa al check-in')
add_bullet(doc, 'Categorías: alimentos, bebidas, higiene, snacks, cuidado personal')
add_bullet(doc, '', bold_prefix='Ticket promedio estimado: $18–$35')
add_bullet(doc, '', bold_prefix='Tasa de adopción esperada: 35% de huéspedes')
add_bullet(doc, '', bold_prefix='Margen bruto promedio: 42%')

add_sub_subsection(doc, 'Máquinas Expendedoras')
add_bullet(doc, 'Unidades con pantalla táctil y múltiples métodos de pago')
add_bullet(doc, 'Ubicación: edificios residenciales con usuarios Boogie')
add_bullet(doc, '', bold_prefix='Ingreso estimado por unidad: $400–$800/mes')
add_bullet(doc, '', bold_prefix='Costo de adquisición por unidad: $3,500–$5,000')
add_bullet(doc, '', bold_prefix='Período de recuperación: 6–10 meses')

# 2.4
add_horizontal_line(doc, '52B788')
add_subsection_title(doc, 'Fuente 2.4 — Alquiler de Servicios de Terceros')

add_body(doc, 'Durante la estancia, el huésped puede contratar servicios complementarios a través de la plataforma: consolas de videojuegos, chef ejecutivo, maquillador, peluquero, servicio de agua potable, fotografía profesional, transporte y más. Boogie retiene una comisión del 15% al 25% por cada servicio reservado.')

add_body(doc, 'Este modelo presenta cero costo de inventario y escala proporcionalmente con el volumen de huéspedes activos. Los proveedores acceden a una demanda cautiva; los huéspedes obtienen conveniencia. Boogie monetiza la intermediación sin riesgo operativo.')

add_bullet(doc, '', bold_prefix='Comisión promedio: 22%')
add_bullet(doc, '', bold_prefix='Margen bruto: 95%')
add_bullet(doc, '', bold_prefix='Ingreso est. mes 6: $800')

# 2.5
add_horizontal_line(doc, '52B788')
add_subsection_title(doc, 'Fuente 2.5 — Flipping Inmobiliario')
add_body(doc, 'Véase detalle completo en la Sección 04.')

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECCIÓN 04 — FLIPPING INMOBILIARIO
# ═══════════════════════════════════════════════
add_section_title(doc, 'SECCIÓN 04 — FLIPPING INMOBILIARIO Y ESTRATEGIA DE ACTIVOS')

add_body(doc, 'El modelo de flipping inmobiliario constituye el componente central de la estrategia de largo plazo de Boogie. Mediante la adquisición progresiva de inmuebles, su reforma integral y posterior posicionamiento en el mercado de alquileres vacacionales, Boogie forma una cartera de bienes que se rotan constantemente, obteniendo porcentajes de rentabilidad gracias al gentrificado del bien, su posicionamiento dentro del mercado y la revalorización propia del metro cuadrado producto de factores macroeconómicos nacionales.')

add_subsection_title(doc, 'Mecanismo de Generación de Valor')

add_sub_subsection(doc, 'Gentrificación del activo')
add_body(doc, 'Reforma integral que incrementa el valor de mercado del inmueble. Transformación de propiedad estándar a espacio premium. Margen de flipping estimado: 40%–80%.')

add_sub_subsection(doc, 'Revalorización macroeconómica')
add_body(doc, 'El metro cuadrado en zonas estratégicas se revaloriza por factores estructurales: retorno de diáspora, inversión extranjera y reactivación comercial.')

add_sub_subsection(doc, 'Ingresos por ocupación')
add_body(doc, 'Mientras el inmueble permanece en cartera, genera ingresos por alquiler vacacional a través de la plataforma. El activo produce rendimiento durante el período de tenencia.')

add_spacer(doc, 4)
add_body(doc, 'El ciclo operativo contempla una rotación de 12 a 18 meses por activo: adquisición, reforma, posicionamiento en plataforma, generación de ingresos por ocupación, y venta con prima de revalorización. Cada ciclo alimenta el siguiente, creando un efecto compuesto sobre el patrimonio.')

add_subsection_title(doc, 'Proyección del Inmueble Piloto')

add_table(doc,
    ['Indicador', 'Conservador', 'Optimista'],
    [
        ['Costo de adquisición', '$18,000', '$25,000'],
        ['Costo de reforma', '$5,000', '$8,000'],
        ['Valor post-reforma', '$32,000', '$48,000'],
        ['Ingreso mensual por alquiler', '$800/mes', '$1,400/mes'],
        ['Margen de flipping (en caso de venta)', '+40%', '+80%'],
        ['ROI anual compuesto', '120%', '200%+'],
    ]
)

add_body(doc, 'Ventaja del modelo integrado: A diferencia de un operador de flipping tradicional, Boogie puede generar ingresos por ocupación mientras espera el momento óptimo de venta. Cada inmueble en la plataforma genera tráfico, reseñas y reputación que alimentan el ecosistema completo.')

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECCIÓN 05 — RUTA ESTRATÉGICA
# ═══════════════════════════════════════════════
add_section_title(doc, 'SECCIÓN 05 — RUTA ESTRATÉGICA')

add_body(doc, 'Hoja de ruta operativa y financiera a 24 meses, estructurada en fases con hitos medibles y objetivos cuantificables.')

phases = [
    ('MESES 1–3 — FASE DE FUNDACIÓN',
     'Lanzamiento de la plataforma en versión MVP. Adquisición del inmueble piloto e inicio de proceso de reforma. Despliegue de las dos primeras máquinas expendedoras en residenciales aliados. Onboarding de los primeros 50 anfitriones y 20 propiedades activas. Inicio de campaña de marketing digital.'),
    ('MESES 4–6 — FASE DE TRACCIÓN',
     'Superior a 200 propiedades listadas en la plataforma. Generación de primeros ingresos recurrentes por comisiones y suscripciones. Inmueble piloto operativo generando ingresos por ocupación. Contratos firmados con al menos 4 residenciales para máquinas expendedoras.'),
    ('MESES 7–12 — FASE DE ESCALAMIENTO',
     '500+ propiedades listadas. ARR de suscripciones superior a $77,000. Lanzamiento del servicio "Arma Tu Boogie". Inicio del proceso de adquisición del segundo inmueble. Despliegue de 5 máquinas expendedoras operativas.'),
    ('MESES 13–18 — FASE DE CONSOLIDACIÓN',
     'Primer ciclo de flipping completado con retorno verificado. Cartera de 3 a 4 inmuebles en operación. Más de 1,000 propiedades en plataforma. Expansión a segunda ubicación geográfica. Ingresos mensuales superando $30,000.'),
    ('MESES 19–24 — FASE DE POSICIONAMIENTO',
     'Posición consolidada como referente del mercado vacacional venezolano. Cartera de 6 a 8 inmuebles. 15 o más máquinas expendedoras operativas. Evaluación de opciones para ronda de inversión Serie A o participación accionaria estratégica.'),
]

for title, desc in phases:
    add_sub_subsection(doc, title)
    add_body(doc, desc)

add_horizontal_line(doc, '52B788')
add_subsection_title(doc, 'Filosofía de Reinversión')

add_body(doc, 'Las utilidades generadas por la plataforma se destinan en su totalidad a la capitalización de los dos modelos de negocio principales: inversión inmobiliaria y venta de commodities. Esta política asegura que el crecimiento de la empresa esté respaldado por activos tangibles, no únicamente por flujo de caja operativo.')

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECCIÓN 06 — PROYECCIONES FINANCIERAS
# ═══════════════════════════════════════════════
add_section_title(doc, 'SECCIÓN 06 — PROYECCIONES FINANCIERAS')

add_body(doc, 'Proyecciones basadas en escenarios conservadores, utilizando benchmarks del mercado de alquiler de corta estancia en Latinoamérica ajustados a las condiciones del mercado venezolano.')

add_subsection_title(doc, 'Ingresos por Línea de Negocio — Escenario Conservador')

add_table(doc,
    ['Línea de Negocio', 'Mes 6', 'Mes 12', 'Mes 18', 'Mes 24'],
    [
        ['Comisiones por intermediación', '$2,800', '$6,500', '$12,000', '$18,000'],
        ['Suscripciones de anfitriones', '$2,100', '$6,480', '$10,800', '$16,200'],
        ['Venta de commodities', '$1,200', '$4,800', '$9,600', '$15,000'],
        ['Servicios de terceros', '$800', '$2,400', '$5,200', '$8,500'],
        ['Ingresos inmobiliarios', '$0', '$800', '$3,200', '$6,400'],
        ['TOTAL MENSUAL', '$6,900', '$20,980', '$40,800', '$64,100'],
    ]
)

add_subsection_title(doc, 'Indicadores Clave de Desempeño — Mes 1 a Mes 24')

add_table(doc,
    ['KPI', 'Inicio', 'Final', 'Crecimiento'],
    [
        ['Propiedades listadas', '50', '1,200', '2,300%'],
        ['Anfitriones con suscripción activa', '15', '200', '1,233%'],
        ['Máquinas expendedoras desplegadas', '2', '15', '650%'],
        ['Inmuebles en cartera', '1', '8', '700%'],
        ['Valor total cartera inmobiliaria', '$23,000', '$280,000+', '1,117%'],
    ]
)

add_subsection_title(doc, 'Punto de Equilibrio')

add_body(doc, 'Bajo el escenario conservador, Boogie alcanza el break-even operativo entre los meses 14 y 18. A partir de ese punto, la totalidad de las utilidades se reinvierte en adquisición de inmuebles y expansión de la red de commodities, acelerando el crecimiento del patrimonio corporativo.')

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECCIÓN 07 — OBJETIVOS DEL FINANCIAMIENTO
# ═══════════════════════════════════════════════
add_section_title(doc, 'SECCIÓN 07 — OBJETIVOS DEL FINANCIAMIENTO')

add_body(doc, 'Cada componente del financiamiento tiene una asignación específica con retorno cuantificable. No se contempla un burn rate abstracto: se financian activos, contratos y usuarios concretos.')

funding_items = [
    ('01. Adquisición de Inmueble Piloto ($18,000–$25,000)',
     'Un inmueble que fungirá como modelo operativo del negocio. Genera ingresos por alquiler vacacional desde el primer mes y se revaloriza como activo tangible. Constituye la prueba de concepto del modelo de flipping y garantía de solvencia para el inversionista.'),
    ('02. Adquisición de Máquinas Expendedoras ($7,000–$10,000)',
     'Dos unidades de última generación para firmar de manera inmediata dos contratos con dos entidades residenciales distintas, estableciendo los primeros dos aliados comerciales ejecutivos de Boogie.'),
    ('03. Campaña de Adquisición de Usuarios ($5,000–$10,000)',
     'Inversión exclusiva en publicidad digital y pago a creadores de contenido. Boogie cuenta con equipo de producción audiovisual interno, por lo que los costos de producción no aplican. La totalidad del presupuesto se destina a distribución y alcance.'),
    ('04. Reserva Operativa ($5,000–$10,000)',
     'Fondo destinado a operación, contingencias y aprovechamiento de oportunidades de mercado que puedan presentarse durante los primeros meses.'),
]

for title, desc in funding_items:
    add_sub_subsection(doc, title)
    add_body(doc, desc)

add_subsection_title(doc, 'Desglose por Escenarios')

add_table(doc,
    ['Concepto', 'Mínimo', 'Óptimo', 'Máximo'],
    [
        ['Inmueble piloto', '$18,000', '$22,000', '$25,000'],
        ['Máquinas expendedoras (×2)', '$7,000', '$8,500', '$10,000'],
        ['Marketing y adquisición', '$5,000', '$7,500', '$10,000'],
        ['Reserva operativa', '$5,000', '$7,000', '$10,000'],
        ['TOTAL FINANCIAMIENTO', '$35,000', '$45,000', '$55,000'],
    ]
)

add_body(doc, 'Entre el 55% y el 60% del financiamiento se convierte directamente en un activo inmobiliario con valor de mercado superior al costo de adquisición. El financiamiento no sustenta un concepto: adquiere un activo.')

doc.add_page_break()

# ═══════════════════════════════════════════════
# SECCIÓN 08 — PROPUESTA AL SOCIO BANCARIO
# ═══════════════════════════════════════════════
add_section_title(doc, 'SECCIÓN 08 — PROPUESTA AL SOCIO BANCARIO')

add_body(doc, 'Boogie presenta una relación de largo plazo fundamentada en beneficios mutuos, transparencia operativa y alineamiento de intereses estratégicos.')

add_subsection_title(doc, 'Beneficios para la Entidad Financiera')

partner_benefits = [
    ('01. Plataforma sólida con alto potencial de rentabilidad.',
     'Cinco fuentes de ingresos diversificadas, modelo de negocio con respaldo en activos tangibles y proyecciones financieras conservadoras que muestran viabilidad clara.'),
    ('02. Alto tráfico monetario en cuentas oficiales.',
     'En el contexto de apertura económica y posible activación de pagos en moneda extranjera, Boogie garantiza flujo constante de divisas y moneda nacional en las cuentas de la empresa, fortaleciendo la relación comercial con la entidad financiera desde el primer trimestre de operación.'),
    ('03. Oportunidad de inversión accionaria.',
     'En un mercado donde las fintechs venezolanas muestran crecimiento sostenido, y la reactivación social atrae retornados e inversión extranjera, la entidad financiera tiene la posibilidad de adquirir participación accionaria en Boogie a valoración preferencial, con opción de planificar una Oferta Pública de Venta (OPV) cuando la empresa lo considere oportuno.'),
    ('04. Socio comercial de largo plazo.',
     'Boogie se compromete a recompensar la confianza depositada mediante generación consistente de oportunidades de rentabilidad y lealtad sostenida en el tiempo.'),
]

for title, desc in partner_benefits:
    add_sub_subsection(doc, title)
    add_body(doc, desc)

add_subsection_title(doc, 'Modalidades de Relación Propuestas')

add_sub_subsection(doc, 'Opción A — Préstamo con Garantía')
add_bullet(doc, 'Plazo de 18 a 24 meses')
add_bullet(doc, 'Tasa preferencial por volumen de operaciones')
add_bullet(doc, 'Garantía: inmueble adquirido más flujo de caja')
add_bullet(doc, 'Inicio de pagos mensuales a partir del mes 6')
add_bullet(doc, 'Relación bancaria exclusiva para operaciones de la empresa')

add_sub_subsection(doc, 'Opción B — Inversión con Participación Accionaria')
add_bullet(doc, 'Participación del 5% al 15% del capital')
add_bullet(doc, 'Boogie mantiene control operativo')
add_bullet(doc, 'Dividendos proporcionales a partir del mes 12')
add_bullet(doc, 'Derecho de primera opción en futuras rondas de inversión')
add_bullet(doc, 'Asiento en consejo consultivo')

add_subsection_title(doc, 'Próximos Pasos')

add_body(doc, 'Boogie está disponible para una reunión presencial en la que se presentará la plataforma en funcionamiento, se detallarán los estados financieros proyectados y se discutirán las condiciones específicas de la relación.')

add_spacer(doc, 24)
add_horizontal_line(doc, '1B4332')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Boogie — Documento Confidencial — 2025')
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = MID_TEXT
run.font.name = FONT_BODY
set_paragraph_spacing(p, before=4, after=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Contacto: contacto@boogie.com — @boogie.ve')
run.font.size = Pt(9)
run.font.color.rgb = BRAND_PRIMARY
run.font.name = FONT_BODY
set_paragraph_spacing(p, before=0, after=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documento con clasificación de confidencialidad')
run.font.size = Pt(8)
run.font.italic = True
run.font.color.rgb = LIGHT_TEXT
run.font.name = FONT_BODY
set_paragraph_spacing(p, before=0, after=0)

doc.save(OUTPUT)
print(f'Documento guardado en: {OUTPUT}')
